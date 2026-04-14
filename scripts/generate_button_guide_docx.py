"""
Generate a Turkish Word (.docx) user guide that explains every button in
the Patent Drafting Tool UI, screen by screen.

Run from the repo root:
    python scripts/generate_button_guide_docx.py

Output: docs/Patent_Drafting_Tool_Button_Guide.docx
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parent.parent / "docs" / "Patent_Drafting_Tool_Button_Guide.docx"

SCREENSHOTS_DIR = Path(
    r"C:\Users\melike\OneDrive - CC Bilgi Teknolojileri\Resimler\Ekran görüntüleri"
)
SCREENSHOTS = {
    "dashboard": SCREENSHOTS_DIR / "Ekran görüntüsü 2026-04-14 011914.png",
    "workspace_full": SCREENSHOTS_DIR / "Ekran görüntüsü 2026-04-14 011929.png",
    "workspace_alt": SCREENSHOTS_DIR / "Ekran görüntüsü 2026-04-14 012018.png",
    "draft_editor": SCREENSHOTS_DIR / "Ekran görüntüsü 2026-04-14 012112.png",
    "upload_excel": SCREENSHOTS_DIR / "Ekran görüntüsü 2026-04-14 012052.png",
    "extract_modal": SCREENSHOTS_DIR / "Ekran görüntüsü 2026-04-14 012036.png",
    "project_settings": SCREENSHOTS_DIR / "Ekran görüntüsü 2026-04-14 012311.png",
}

doc = Document()

# Page margins
for section in doc.sections:
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

# Default style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def add_cover():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(6):
        p.add_run().add_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Patent Drafting Tool")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x1E, 0x3C, 0x82)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Buton Rehberi")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x1E, 0x3C, 0x82)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Arayüzdeki her butonun ne işe yaradığını ve")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("nasıl kullanılacağını anlatan kullanıcı kılavuzu")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Bu rehber, Patent Drafting Tool arayüzünün ana ekranlarındaki "
        "ve modal pencerelerindeki tüm butonları tek tek açıklar. Her "
        "buton için ne işe yaradığı, nasıl kullanıldığı ve dikkat "
        "edilmesi gereken durumlar verilmiştir."
    )
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Sürüm: feature/inventor-qa-section")
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()


def section_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Background shading via XML
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "1E3C82")
    p_pr = p._p.get_or_add_pPr()
    p_pr.append(shd)


def sub_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x1E, 0x3C, 0x82)


def intro(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_screenshot(key, width_inches=6.2, caption=None):
    path = SCREENSHOTS.get(key)
    if not path or not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(8)
        cr = cp.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(9)
        cr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(10)


def kv_bullet(label, body):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + " ")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(body)
    r2.font.size = Pt(10)


def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def button_block(name, what, how, note=None):
    # Button name
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("• " + name)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x10, 0x10, 0x10)

    # What
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Ne işe yarar: ")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(what)
    r2.font.size = Pt(10)

    # How
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Nasıl kullanılır: ")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(how)
    r2.font.size = Pt(10)

    # Note
    if note:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run("Not: ")
        r.bold = True
        r.italic = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xB0, 0x60, 0x00)
        r2 = p.add_run(note)
        r2.italic = True
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0xB0, 0x60, 0x00)
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ─────────────────────────────────────────────────────────────────
# COVER
# ─────────────────────────────────────────────────────────────────
add_cover()


# ─────────────────────────────────────────────────────────────────
# 1) ANA SAYFA — DASHBOARD
# ─────────────────────────────────────────────────────────────────
section_heading("1. Ana Sayfa (Dashboard)")
add_screenshot("dashboard", caption="Ana sayfa: sol panelde proje listesi, sağda tanıtım panosu")
intro(
    "Uygulamayı açtığınızda gördüğünüz ilk ekran. Sol tarafta proje "
    "(workspace) listesi, sağ tarafta ise uygulamanın tanıtım panosu yer "
    "alır. Üstte ortak bir navbar bulunur."
)

sub_heading("Üst Navbar (Ana Sayfada)")
button_block(
    "PATENT DRAFTING TOOL (logo + başlık)",
    "Uygulamanın markasını gösterir. Tıklanabilir bir aksiyonu yoktur.",
    "Sadece görsel — başka bir sekmedeyseniz ana sayfaya dönmek için Back to Home butonunu kullanın.",
)
button_block(
    "Data Loaded / No Data (durum rozeti)",
    "Excel verisinin (RAG vektör veritabanı) yüklü olup olmadığını gösterir. "
    "Yeşil 'Data Loaded' = embedding'ler hazır; gri 'No Data' = henüz Excel yüklenmedi.",
    "Sadece bilgi amaçlıdır. AI Suggest Definition gibi özellikler için rozetin yeşil olması gerekir.",
)
button_block(
    "Back to Home",
    "Açık olan workspace'ten çıkıp ana sayfaya (proje listesine) döner.",
    "Workspace'teyken bu butona tıklayın; seçili proje sıfırlanır ve dashboard ekranı gelir.",
)
button_block(
    "Upload Data",
    "Excel veri yükleme modalını açar (Upload Excel Data). Aynı işlevi soldaki 'Upload Excel Data' butonu da görür.",
    "Tıklayın → çıkan modaldan .xlsx dosyasını seçin → embedding'ler arka planda oluşturulur.",
)

sub_heading("Sol Panel — Projects")
button_block(
    "Upload Excel Data",
    "Sol panelden Excel yükleme modalını açar. Üst navbardaki Upload Data ile aynı şeyi yapar; kolaylık için iki yerde de bulunur.",
    "Tıklayın → modal açılır → Excel dosyasını seçip yükleyin.",
)
button_block(
    "+ New Project",
    "Yeni bir patent projesi oluşturmak için 'New Project' modalını açar.",
    "Tıklayın → Project Name (proje adı) ve Patent Owner (sahibi) alanlarını doldurun → Create butonuna basın. Proje listesine eklenir.",
)

sub_heading("Proje Kartı (her proje için)")
intro(
    "Listedeki her projenin altında üç buton bulunur. Kartın kendisine "
    "tıklamak da projeyi açar."
)
button_block(
    "Inputs",
    "Projeye ait Patent Inputs modalını açar. Şu anda bu modalın içinde 'Buluşçu ile Yazışmalar (Inventor_QA)' kartı vardır: Q&A metnini yazabilir ve doküman ekleyebilirsiniz.",
    "Projenin yanındaki Inputs butonuna tıklayın → metin yazıp Save Q&A Text ile kaydedin veya Choose document ile dosya seçip Upload ile yükleyin. Mevcut dokümanlar listelenir, indirilebilir veya silinebilir.",
)
button_block(
    "Edit",
    "Projenin adını ve patent sahibini düzenlemek için Edit Project modalını açar.",
    "Edit'e tıklayın → Project Name ve/veya Patent Owner alanlarını değiştirin → Save ile kaydedin.",
)
button_block(
    "Open",
    "Projeyi açar ve workspace ekranına (Claims, Elements, Draft Editor) geçer. Karta herhangi bir yerine tıklamak da aynı işi yapar.",
    "Tıklayın → seçili proje workspace'te yüklenir.",
)


# ─────────────────────────────────────────────────────────────────
# 2) WORKSPACE — ÜST NAVBAR
# ─────────────────────────────────────────────────────────────────
doc.add_page_break()
section_heading("2. Workspace — Üst Navbar")
add_screenshot("workspace_full", caption="Workspace genel görünüm: 3 panel düzeni ve üst navbar")
intro(
    "Bir proje açıldığında üstteki navbar, projeyle ilgili bilgileri "
    "gösterir ve hızlı aksiyonlar sunar."
)
button_block(
    "Proje Adı (örn: deneme1)",
    "Açık olan projenin adını gösterir. Yanında küçük gri yazıyla proje sahibi (— melike) görünür.",
    "Sadece bilgi amaçlıdır; düzenlemek için yanındaki Rename butonunu kullanın.",
)
button_block(
    "Rename",
    "Açık projenin adını/sahibini değiştirmek için Edit Project modalını açar (proje listesindeki Edit ile aynı modal).",
    "Tıklayın → yeni adı ve/veya owner'ı yazın → Save'e basın. Navbar ve proje listesi otomatik güncellenir.",
)
button_block(
    "Back to Home",
    "Workspace'ten çıkıp ana sayfaya (proje listesi) döner.",
    "Tıklayın; üzerinde çalıştığınız değişiklikler kaydedildiyse korunur.",
)
button_block(
    "Upload Data",
    "Excel veri yükleme modalını açar (Workspace içinden de Excel yüklenebilir).",
    "Tıklayın → Excel dosyasını seçip yükleyin.",
)


# ─────────────────────────────────────────────────────────────────
# 3) WORKSPACE — SOL PANEL: CLAIMS & ELEMENTS
# ─────────────────────────────────────────────────────────────────
section_heading("3. Workspace — Sol Panel: Claims & Elements")
add_screenshot("workspace_alt", caption="Sol panelde Claims & Elements; element ağacı, ekleme/silme ve sıralama kontrolleri")
intro(
    "Patentin tüm istemlerini (claims) ve her istemin element ağacını "
    "yönettiğiniz alandır. Üstte 'Claim Structures' başlığı yer alır."
)
button_block(
    "+ Add Claim",
    "Yeni bir istem (claim) eklemek için 'Add Claim' modalını açar.",
    "Tıklayın → Dependency Type (Independent/Dependent) ve Category (Apparatus/Method) seçin; Dependent seçtiyseniz Parent Claim belirleyin → Add Claim ile oluşturun.",
)

sub_heading("Her bir Claim kartında")
button_block(
    "Claim Numarası (Claim 1, Claim 2, ...)",
    "İstemi tanımlayan başlık. Yanında durum rozeti (Ready / Incomplete) bulunabilir: 'Ready' = bağlı tüm element'lerin tanımı tamam; 'Incomplete' = en az birinin tanımı eksik.",
    "Karta tıklayarak istemi seçin; sağdaki Claim Draft Editor o istemin metnini gösterir.",
)
button_block(
    "Genişlet/Daralt oku",
    "Claim kartının altındaki ELEMENT TREE bölümünü açar veya kapatır.",
    "Ok simgesine tıklayın; element ağacı görünür/gizlenir.",
)
button_block(
    "+ Add Elements",
    "Patent havuzundaki bir element'i bu istem'e bağlamak için Link Element modalını açar.",
    "Tıklayın → açılan modaldan element seçin → Link ile bağlayın. Element ağacında yeni satır olarak görünür.",
)
button_block(
    "Çöp ikonu (claim sil)",
    "Claim kartının sağ üstündeki çöp ikonu; bu istemi siler.",
    "Tıklayın → onay kutusu çıkar → Delete'e basarsanız claim silinir. Geri alınamaz.",
    note="Silmek istemediğiniz bir istem için tıklamadığınızdan emin olun.",
)

sub_heading("Element Tree (her claim altında)")
button_block(
    "Yukarı/Aşağı oku (sıralama)",
    "Seçili element'in claim içindeki sırasını bir yukarı veya bir aşağı taşır.",
    "Önce element'in üzerine tıklayarak SELECTED rozetini almasını sağlayın → yukarı/aşağı oklarıyla sırayı değiştirin.",
)
button_block(
    "Yenile (refresh)",
    "Bu claim'in element ağacını backend'den yeniden yükler.",
    "Sıralama veya bağlantılarda görünmeyen bir değişiklik olursa tıklayın.",
)
button_block(
    "Search elements (arama kutusu)",
    "Element ağacını isimle filtreler.",
    "Aramak istediğiniz kelimeyi yazın → eşleşmeyen element'ler gizlenir.",
)
button_block(
    "Element satırı",
    "Element'in adını gösterir. Sağda durum ikonu (tik = tanım tamam / uyarı = tanım eksik) vardır.",
    "Tıklayın → element seçilir ve Element Definition modalı açılır.",
)
button_block(
    "Çöp ikonu (element bağlantısını çöz)",
    "Bu element'i bu claim'den ayırır (element silinmez, sadece bağ kalkar).",
    "Tıklayın; element claim'in ağacından çıkar ama patent havuzunda kalmaya devam eder.",
)


# ─────────────────────────────────────────────────────────────────
# 4) WORKSPACE — ORTA PANEL: LLM EXTRACTION QUEUE
# ─────────────────────────────────────────────────────────────────
doc.add_page_break()
section_heading("4. Workspace — Orta Panel: LLM Extraction Queue")
add_screenshot("workspace_full", caption="Orta panel: patent element havuzu (LLM Extraction Queue)")
intro(
    "Patent havuzundaki tüm element'lerin listelendiği panel. Buradaki "
    "element'ler claim'lere bağlanabilir ve tanımları düzenlenebilir."
)
button_block(
    "Extract",
    "BBF + Report dökümanlarından otomatik element çıkarmak için 'Extract Elements from BBF + Report' modalını açar. Bu modalda ek olarak Buluşçu ile Yazışmalar (Inventor_QA) dokümanı da yüklenebilir.",
    "Tıklayın → BBF, Report (zorunlu) ve isterseniz Inventor_QA dosyalarını seçin → Extract Elements butonuna basın. İşlem bitince çıkarılan element'ler havuza eklenir.",
)
button_block(
    "+ Add Element",
    "Boş bir Element Definition modalı açar; havuza manuel olarak yeni element eklemenizi sağlar.",
    "Tıklayın → Element Name (zorunlu), Reference Number ve Definition alanlarını doldurun → Save & Close ile kaydedin.",
)

sub_heading("Element listesindeki her satırda")
button_block(
    "DRAG (sürükleme tutamağı)",
    "Element'i sürükleyip claim kartlarının üzerine bırakarak hızlıca bağlama imkânı verir.",
    "DRAG yazısının üzerine basılı tutun → istediğiniz claim kartının üzerine bırakın → element o claim'e link'lenir.",
)
button_block(
    "Edit",
    "Element Definition modalını düzenleme modunda açar.",
    "Tıklayın → element adı, reference number, definition, AI suggest, linked claims listesi gibi tüm alanlara erişebilirsiniz.",
)
button_block(
    "Çöp ikonu (element sil)",
    "Element'i havuzdan tamamen siler. Bağlı olduğu tüm claim'lerden de düşer.",
    "Onay verdikten sonra silinir; geri alınamaz.",
    note="Element kalıcı olarak silinir; sadece bağlantıyı çözmek istiyorsanız claim altındaki çöp ikonunu (unlink) kullanın.",
)


# ─────────────────────────────────────────────────────────────────
# 5) WORKSPACE — SAĞ PANEL: CLAIM DRAFT EDITOR
# ─────────────────────────────────────────────────────────────────
doc.add_page_break()
section_heading("5. Workspace — Sağ Panel: Claim Draft Editor")
add_screenshot("draft_editor", caption="Claim Draft Editor: seçili istem için taslak metin alanı ve aksiyon butonları")
intro(
    "Seçili istem'in (claim) draft (taslak) metnini yazıp düzenlediğiniz "
    "alan. Sol panelden bir claim seçtiğinizde buradaki textarea aktifleşir."
)
button_block(
    "Settings",
    "Project Settings modalını açar; Invention Context, BBF Text ve LLM API URL alanlarını içerir.",
    "Tıklayın → Invention Context (örn: 'Antenna assembly for an aircraft'), BBF Text (BBF'in metin içeriği) ve isteğe bağlı Colab ngrok LLM URL'sini girin → Save'e basın.",
)
button_block(
    "Assemble with Report",
    "Tüm tanımlanmış element'lerden bir 'ELEMENT DEFINITIONS REPORT' üretip draft alanına yazar.",
    "Tıklayın → element listesinden tanımı dolu olanlar bir araya getirilip düzenli bir rapor olarak draft textarea'sına yazılır.",
    note="Mevcut draft metninin üzerine yazar; önce kaydetmek istiyorsanız Save'e basın.",
)
button_block(
    "Save",
    "Seçili claim'in draft metnini PostgreSQL veritabanına kaydeder. Sonra kısa bir 'Saved' geri bildirimi gösterir.",
    "Önce sol panelden bir claim seçin → metni düzenleyin → Save'e basın. Buton kısaca 'Saving…' → 'Saved' olur.",
    note="Hiçbir claim seçili değilse uyarı verir.",
)
button_block(
    "Go to AI Draft",
    "Tüm claim'lerin yapılandırılmış metnini Draft Composer ekranına aktarıp AI ile tam patent draft'ı üretmek için yönlendirir.",
    "Tıklayın → composer ekranına geçersiniz; oradan Generate Draft ile AI üretimi başlatabilirsiniz.",
)
button_block(
    "Draft Composer",
    "Doğrudan Draft Composer (AI Draft Workspace) sayfasına geçiş yapar.",
    "Tıklayın → Structured Claims paneline mevcut claim'leri girip Draft Output panelinde AI çıktısı üretebilirsiniz.",
)


# ─────────────────────────────────────────────────────────────────
# 6) MODAL — UPLOAD EXCEL DATA
# ─────────────────────────────────────────────────────────────────
doc.add_page_break()
section_heading("6. Modal — Upload Excel Data")
add_screenshot("upload_excel", caption="Upload Excel Data modalı — embedding'ler hazır olduğunda yeşil onay mesajı")
intro(
    "RAG (Retrieval-Augmented Generation) için kullanılan vektör "
    "veritabanını oluşturan Excel dosyasını yükleme ekranı."
)
button_block(
    "Click to select Excel file (.xlsx)",
    "Bilgisayarınızdan bir .xlsx dosyası seçmek için sistem dosya tarayıcısını açar.",
    "Tıklayın → çevirisi yapılmış Excel dosyasını seçin. Yükleme ve embedding oluşturma otomatik başlar; durumu altta görürsünüz (örn: 'Loaded 9133 docs. FAISS index ready.').",
    note="Embedding oluşturma birkaç dakika sürebilir; modalı kapatsanız bile arka planda tamamlanır.",
)
button_block(
    "Close (× / Close)",
    "Modalı kapatır.",
    "Yükleme bittiyse veya devam ediyorsa kapatabilirsiniz; arka plan işlemi etkilenmez.",
)


# ─────────────────────────────────────────────────────────────────
# 7) MODAL — EXTRACT ELEMENTS FROM BBF + REPORT
# ─────────────────────────────────────────────────────────────────
section_heading("7. Modal — Extract Elements from BBF + Report")
add_screenshot("extract_modal", caption="Extract Elements modalı: BBF, Report ve Buluşçu ile Yazışmalar (Inventor_QA) yükleme alanları")
intro(
    "BBF (Buluş Bildirim Formu) ve Research Report dökümanlarını yükleyerek "
    "otomatik element çıkarımı yapan modal. Ayrıca isteğe bağlı 'Buluşçu "
    "ile Yazışmalar (Inventor_QA)' dosyası da yükleyebilirsiniz."
)
button_block(
    "BBF Document — Click to select BBF file",
    "BBF (Buluş Bildirim Formu) dökümanını seçmek için kullanılır. Pipeline bu dosyadan unsurları (elements) çıkarır.",
    "Tıklayın → .docx, .pdf veya .txt formatında BBF dosyanızı seçin.",
)
button_block(
    "Report Document — Click to select Report file",
    "Araştırma raporu (Research Report) dökümanını seçer; pipeline BBF ile birlikte bu dosyayı kullanarak element ve tanımları çıkarır.",
    "Tıklayın → .docx, .pdf veya .txt formatında raporunuzu seçin.",
)
button_block(
    "Buluşçu ile Yazışmalar (Inventor_QA) — Click to select Inventor_QA file",
    "Buluşçu ile yapılan yazışma/açıklama dökümanlarını yükler. Bu alan opsiyoneldir. Dosya seçer seçmez otomatik olarak ilgili patent'in inventor_qa kayıtlarına yüklenir ve modalı kapatsanız dahi listede kalır.",
    "Tıklayın → .pdf, .docx, .txt vb. dosyanızı seçin. Yükleme tamam olunca 'Uploaded ...' mesajı görünür ve mevcut dokümanlar altta listelenir.",
    note="Bu dosya pipeline'a verilmez; sadece projenin Inventor_QA dökümanlarına eklenir. Pipeline yalnızca BBF + Report ile çalışır.",
)
button_block(
    "Cancel",
    "Modalı kapatır; çalışmaya başlamış pipeline'ı durdurmaz.",
    "Tıklayın; modal kapanır.",
)
button_block(
    "Extract Elements",
    "Seçili BBF + Report dökümanlarıyla extraction pipeline'ını çalıştırır. Pipeline bittiğinde çıkarılan element'ler havuza (LLM Extraction Queue) otomatik eklenir.",
    "Önce BBF ve Report dosyalarını seçin → tıklayın → 'Running pipeline…' mesajını görürsünüz → tamamlanınca 'Extracted N elements!' yazısı çıkar.",
    note="BBF ve Report her ikisi de zorunludur; biri eksikse uyarı verir.",
)


# ─────────────────────────────────────────────────────────────────
# 8) MODAL — PROJECT SETTINGS
# ─────────────────────────────────────────────────────────────────
doc.add_page_break()
section_heading("8. Modal — Project Settings")
add_screenshot("project_settings", caption="Project Settings modalı: invention context, BBF text ve LLM URL ayarları", width_inches=4.5)
intro(
    "AI üretiminde kullanılacak invention context, BBF text ve LLM ayarlarını "
    "giriş alanı."
)
button_block(
    "Invention Context",
    "Patent için kısa bağlam metni; AI promptlarına context olarak verilir (örn: 'Antenna assembly for an aircraft').",
    "Buluşunuzu en kısa şekilde tarif eden tek satırlık bir cümle yazın.",
)
button_block(
    "BBF Text",
    "BBF (Buluş Bildirim Formu) içeriğinin tam metni. AI Suggest Definition ve diğer üretim modülleri bu metni kullanır.",
    "BBF dökümanından kopyaladığınız metni textarea'ya yapıştırın (ya da Extract Elements ile pipeline çalıştırırsanız sistem otomatik doldurur).",
)
button_block(
    "LLM API URL (Colab ngrok)",
    "Colab'da çalışan Mistral-7B (veya benzeri) LLM'in ngrok adresini girer. Boş bırakılırsa heuristik fallback kullanılır.",
    "Colab notebook'unuzu çalıştırın → ngrok URL'sini kopyalayın → buraya yapıştırın.",
    note="URL yanlışsa veya ngrok düşmüşse AI suggest çalışmaz; boş bırakırsanız basit fallback devreye girer.",
)
button_block(
    "Cancel",
    "Modalı kaydetmeden kapatır.",
    "Tıklayın; girilen değerler kaybolur (henüz Save'lemediyseniz).",
)
button_block(
    "Save",
    "Girilen Invention Context, BBF Text ve LLM URL ayarlarını uygulamanın bellek/cache'ine kaydeder.",
    "Doldurduktan sonra Save'e basın → modal kapanır → ayarlar sonraki AI üretimlerinde kullanılır.",
)


# ─────────────────────────────────────────────────────────────────
# 9) HIZLI İŞ AKIŞI
# ─────────────────────────────────────────────────────────────────
doc.add_page_break()
section_heading("9. Hızlı İş Akışı (Tipik Kullanım)")
intro(
    "Aşağıdaki adımlar, sıfırdan bir patent draft'ı oluşturmak için "
    "izlemeniz gereken en kısa yolu özetler."
)

steps = [
    ("1. Excel verisini yükleyin",
     "Üst navbardaki 'Upload Data' veya sol panelden 'Upload Excel Data' ile çevirisi yapılmış Excel dosyanızı yükleyin. 'Data Loaded' rozeti yeşile dönmeli."),
    ("2. Yeni proje oluşturun",
     "Sol paneldeki '+ New Project' ile Project Name ve Patent Owner alanlarını girip Create'e basın."),
    ("3. (İsteğe bağlı) Inventor_QA ekleyin",
     "Proje kartındaki 'Inputs' butonu ile Buluşçu ile Yazışmalar metnini yazın ve/veya doküman yükleyin."),
    ("4. Projeyi açın",
     "Proje kartına veya 'Open' butonuna tıklayın; workspace açılır."),
    ("5. Element'leri çıkarın",
     "Orta paneldeki 'Extract' butonuyla BBF + Report (ve isterseniz Inventor_QA) dosyalarını yükleyip pipeline'ı çalıştırın. Çıkan element'ler havuza eklenir."),
    ("6. Claim'leri oluşturun",
     "Sol panelden '+ Add Claim' ile bir veya birden fazla istem ekleyin (independent / dependent + apparatus / method)."),
    ("7. Element'leri claim'lere bağlayın",
     "Her claim altındaki '+ Add Elements' veya orta paneldeki DRAG ile element'leri uygun claim'lere bağlayın."),
    ("8. Tanımları tamamlayın",
     "Element'lerin Edit modalını açıp Definition alanını doldurun. Gerekirse 'AI Suggest Definition' ile öneri alın."),
    ("9. Draft yazın ve kaydedin",
     "Bir claim seçin → sağdaki textarea'ya draft metnini yazın → Save'e basın. 'Assemble with Report' ile element tanımlarını otomatik raporlayabilirsiniz."),
    ("10. AI ile tam draft üretin",
     "'Go to AI Draft' veya 'Draft Composer' ile composer ekranına geçin → 'Generate Draft' ile AI çıktısı üretin."),
]

for h, body in steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x1E, 0x3C, 0x82)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(body)
    r.font.size = Pt(11)


# ─────────────────────────────────────────────────────────────────
# 10) TEKNİK BİLGİ — KOD VE FRONTEND MİMARİSİ
# ─────────────────────────────────────────────────────────────────
doc.add_page_break()
section_heading("10. Teknik Bilgi — Kod ve Frontend Mimarisi")
intro(
    "Hocaya anlatırken kullanabileceğiniz, projenin nasıl yapılandırıldığını "
    "ve hangi teknolojinin nerede çalıştığını gösteren teknik özet."
)

sub_heading("10.1 Kullanılan Teknolojiler")
kv_bullet("Backend framework:", "FastAPI (Python 3.12) — REST API ve statik dosya sunumu için")
kv_bullet("Veritabanı:", "PostgreSQL 16 — projeler, claim'ler, element'ler, inventor_qa ve dökümanlar buraya yazılır")
kv_bullet("ORM:", "SQLAlchemy 2.x (Mapped tipli modeller) + Alembic (şema migration'ları)")
kv_bullet("Vektör veritabanı:", "ChromaDB / FAISS + sentence-transformers (intfloat/multilingual-e5-base)")
kv_bullet("AI / LLM:", "RAG (Retrieval-Augmented Generation) ve Colab üzerinde Mistral-7B (ngrok ile uzaktan)")
kv_bullet("Frontend:", "Vanilla HTML + CSS + JavaScript (build adımı yok). Tek sayfa uygulaması (SPA) gibi davranır.")
kv_bullet("Containerization:", "Docker Compose ile iki servis: PostgreSQL ve FastAPI uygulaması")
kv_bullet("Versiyon kontrolü:", "Git + GitHub (feature branch akışı ile)")

sub_heading("10.2 Klasör Yapısı (önemli yerler)")
code_block(
    "paten_draft_backend/\n"
    "├── app/\n"
    "│   ├── main.py                FastAPI uygulaması ve route bağlama\n"
    "│   ├── database.py            SQLAlchemy session ve Base\n"
    "│   ├── config.py              Ortam ayarları (DB URL, upload path)\n"
    "│   ├── models/                ORM modelleri (Patent, Claim, Element, InventorQA…)\n"
    "│   ├── schemas/               Pydantic istek/cevap şemaları\n"
    "│   ├── services/              İş mantığı (DB üzerinde okuma/yazma)\n"
    "│   ├── routes/                FastAPI router'ları (endpoint'ler)\n"
    "│   ├── ingestion/             Excel okuma + ChromaDB embedding üretimi\n"
    "│   └── rag_engine.py          RAG sorgu motoru\n"
    "├── alembic/versions/          Veritabanı migration dosyaları\n"
    "├── static/\n"
    "│   ├── home.html              Ana sayfa + workspace HTML şablonu\n"
    "│   ├── workspace.html         Alternatif workspace ekranı\n"
    "│   ├── jss/app.js             Tüm frontend mantığı (SPA gibi)\n"
    "│   └── cs/style.css           Koyu tema CSS dosyası\n"
    "├── docs/                      Markdown dokümantasyonu (spec'ler)\n"
    "├── docker-compose.yml         db (postgres) ve app servisleri\n"
    "└── Dockerfile                 Python 3.12 + bağımlılıklar\n"
)

sub_heading("10.3 Frontend Genel Yapısı")
bullet("Tek bir 'home.html' dosyası içinde birden fazla sayfa (page) tanımlıdır: page-dashboard, page-editor, page-composer.")
bullet("Sayfalar arasında geçiş için 'showPage(id)' fonksiyonu kullanılır; class='active' ile yalnızca tek sayfa görünür yapılır.")
bullet("Tüm modal pencereler aynı dosyada tanımlıdır ve 'modal-overlay' class'ına 'active' eklenince açılır.")
bullet("State yönetimi global JavaScript değişkenleriyle yapılır (örn. _projects, _claims, _elements, _patentId).")
bullet("Backend'e istekler 'api(path, opts)' yardımcı fonksiyonu ile atılır (otomatik /api prefix + JSON parse).")
bullet("Cache busting için CSS ve JS dosyalarına versiyon parametresi eklenir (ör. app.js?v=5).")

sub_heading("10.4 Backend Genel Yapısı")
bullet("FastAPI uygulaması 'app/main.py' içinde tanımlanır; routerlar include_router ile bağlanır.")
bullet("Her resource için tipik 4 katman: model (DB tablosu), schema (Pydantic), service (iş mantığı), route (HTTP endpoint).")
bullet("CORS açıktır; statik dosyalar /static prefix ile sunulur, '/' kökü home.html döndürür.")
bullet("Veritabanı oturumu Depends(get_db) ile her isteğe enjekte edilir.")
bullet("Uzun süren işlemler (Excel embedding, BBF pipeline) thread'lerde arka planda çalışır; durum endpoint'i ile poll edilir.")

sub_heading("10.5 Veritabanı Tabloları (ana tablolar)")
kv_bullet("patent:", "patent_id, patent_name, patent_owner, patent_draft, created_at, updated_at")
kv_bullet("claim:", "claim_id, patent_id (FK), claim_number, claim_dependency_type, claim_category, parent_claim_id (FK), claim_text")
kv_bullet("element:", "element_id, patent_id (FK), element_name, reference_number, definition_text")
kv_bullet("claim_element:", "claim_id (FK) + element_id (FK) — many-to-many bağlantı + order_index ile sıralama")
kv_bullet("invention_disclosure:", "patent_id (FK) — prior_art_and_problems, closest_prior_patents, novel_features")
kv_bullet("research_report:", "patent_id (FK) — executive_summary, search_strategy, classification_and_keywords, element_patent_analysis")
kv_bullet("inventor_qa:", "patent_id (FK) — questions_and_answers (text)")
kv_bullet("inventor_qa_document (YENİ):", "qna_id (FK), original_filename, stored_filename, mime_type, size_bytes, created_at")

sub_heading("10.6 Yeni Eklenen 'Buluşçu ile Yazışmalar (Inventor_QA)' Özelliği")
bullet("Amaç: buluşçu ile yapılan görüşme/yazışmaları metin ve doküman olarak proje bazında saklamak.")

p = doc.add_paragraph()
r = p.add_run("Backend katmanında yapılan değişiklikler:")
r.bold = True
r.font.size = Pt(11)
bullet("Model: app/models/inventor_qa.py içine 'InventorQADocument' eklendi (cascade silme ile InventorQA'ya bağlı).")
bullet("Schema: app/schemas/inventor_qa.py oluşturuldu — InventorQARead, InventorQADocumentRead, InventorQAUpdate.")
bullet("Service: app/services/inventor_qa_service.py oluşturuldu — get/upsert text + add/get/list/delete document.")
bullet("Route: app/routes/inventor_qa.py oluşturuldu — /api/patents/{patent_id}/inventor-qa altında 5 endpoint.")
bullet("Migration: alembic/versions/7c0e9b41ad22_add_inventor_qa_document.py — yeni tabloyu oluşturur.")
bullet("Dosyalar diskte settings.uploads_path/inventor_qa/{patent_id}/ altında saklanır; metadata DB'de tutulur.")

p = doc.add_paragraph()
r = p.add_run("Eklenen API endpoint'leri:")
r.bold = True
r.font.size = Pt(11)
code_block(
    "GET    /api/patents/{patent_id}/inventor-qa                  -> Q&A metni + doküman listesi\n"
    "PUT    /api/patents/{patent_id}/inventor-qa                  -> Q&A metnini günceller\n"
    "POST   /api/patents/{patent_id}/inventor-qa/documents        -> doküman yükle\n"
    "GET    /api/patents/{patent_id}/inventor-qa/documents/{id}   -> doküman indir\n"
    "DELETE /api/patents/{patent_id}/inventor-qa/documents/{id}   -> doküman sil"
)

p = doc.add_paragraph()
r = p.add_run("Frontend tarafında yapılan değişiklikler:")
r.bold = True
r.font.size = Pt(11)
bullet("static/home.html: 'Patent Inputs' modalı eklendi — Inventor_QA kartı (textarea + dosya yükleme) içerir.")
bullet("Proje listesindeki her karta '📋 Inputs' butonu eklendi.")
bullet("Workspace navbar'ına '✎ Rename' butonu eklendi (Edit Project modalını açar).")
bullet("Extract Elements modalına 3. yükleme alanı (Buluşçu ile Yazışmalar) eklendi; dosya seçer seçmez otomatik yüklenir.")
bullet("static/jss/app.js: openPatentInputs, loadInventorQa, uploadInventorQaDocument, deleteInventorQaDocument, saveInventorQaText, onInventorQaExtractFileChange fonksiyonları eklendi.")
bullet("static/cs/style.css: patent-input-card, patent-input-doc-row, project-item--stacked stilleri eklendi.")

sub_heading("10.7 Çalıştırma ve Geliştirme Akışı")
code_block(
    "# Tek komutla başlat (Docker Compose ile)\n"
    "docker compose up -d --build\n\n"
    "# Migration'ı container içinde otomatik çalışır (start.sh)\n"
    "# alembic upgrade head\n\n"
    "# Loglar\n"
    "docker compose logs -f app\n\n"
    "# Sağlık kontrolü\n"
    "curl http://localhost:8000/api/health\n\n"
    "# Tarayıcı\n"
    "http://localhost:8000/"
)

sub_heading("10.8 İstek-Cevap Akışı (örnek: Q&A metnini kaydetmek)")
bullet("1. Kullanıcı 📋 Inputs butonuna tıklar → frontend openPatentInputs() çağrılır.")
bullet("2. Modal açılır → loadInventorQa() ile GET /api/patents/{id}/inventor-qa çağrılır → mevcut metin ve dökümanlar gelir.")
bullet("3. Kullanıcı textarea'ya yazıp 'Save Q&A Text' butonuna basar → saveInventorQaText() çağrılır.")
bullet("4. Frontend PUT /api/patents/{id}/inventor-qa endpoint'ine JSON gönderir.")
bullet("5. FastAPI route, service'e devreder; service SQLAlchemy ile UPDATE çalıştırır.")
bullet("6. Cevap olarak güncellenmiş kayıt döner; frontend doküman listesini yeniler.")

sub_heading("10.9 Versiyon Kontrol Akışı")
bullet("master / main → kararlı sürüm")
bullet("feature/inventor-qa-section → bu projedeki yeni özellik branch'i")
bullet("Her değişiklik commit edilir, GitHub'a (melike + origin remote'larına) push edilir.")
bullet("Pull request açılarak code review sonrası master'a merge edilir.")


OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUTPUT))
print(f"Wrote {OUTPUT}")
