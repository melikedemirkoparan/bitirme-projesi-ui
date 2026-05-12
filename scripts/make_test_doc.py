"""Generate a Word document with the Document Assistant test plan.

Run once to produce: C:/Users/melike/Downloads/Document_Assistant_Test_Plani.docx
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


OUTPUT = Path(r"C:\Users\melike\Downloads\Document_Assistant_Test_Plani.docx")


def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
        if color:
            run.font.color.rgb = color
    return h


def add_para(doc, text, bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.name = "Calibri"
        r1.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.name = "Calibri"
        r2.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return p


def add_table_row(table, cells, bold=False, header=False):
    row = table.add_row().cells
    for i, txt in enumerate(cells):
        row[i].text = ""
        p = row[i].paragraphs[0]
        run = p.add_run(txt)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        if bold or header:
            run.bold = True


def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Title
    title = doc.add_heading("Document Assistant — Test Planı", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    add_para(doc, "Patent Drafting Tool — offline_qa_module testi", italic=True, size=11,
             color=RGBColor(0x59, 0x59, 0x59))

    # Ön hazırlık
    add_heading(doc, "Ön hazırlık", level=1)
    add_bullet(doc, "Docker Desktop açık olsun (whale ikonu sabit).")
    add_bullet(doc, "Postgres container çalışıyor olsun: docker ps kontrolü.")
    add_bullet(doc, "Ollama servisi çalışıyor olsun (Windows servisinden otomatik).")
    add_bullet(doc, "qwen2.5:7b modeli indirilmiş olsun: ollama list.")
    add_bullet(doc, "uvicorn ayakta olsun (.\\.venv\\Scripts\\Activate.ps1 + uvicorn app.main:app --reload).")
    add_bullet(doc, "Tarayıcıda localhost:8000 açılıp bir patent Open ile açılmış olsun.")
    add_bullet(doc, "Sağ alt köşede mavi 'Document Assistant' butonuna tıklanmış, drawer açılmış olsun.")

    # Önerilen sıra
    add_heading(doc, "Önerilen test sırası", level=1)
    add_para(doc,
             "Aşağıdaki sırayı izle. Her test bir öncekini doğrular niteliktedir; "
             "P3 en hızlı sonuç verdiği için önce onu deneyip backend bağlantısını "
             "hızlıca onayla, sonra LLM gerektiren P1 ve P2'ye geç.",
             italic=False)

    # TEST 1 — P3
    add_heading(doc, "TEST 1 — Element Lookup (P3)", level=1, color=RGBColor(0xC8, 0x8B, 0x00))
    add_para(doc, "Neden ilk bu: Deterministik metin arama; LLM gerekmez, hızlı döner. "
             "Backend ve DB bağlantısının doğru çalıştığını hemen anlarsın.", italic=True)

    add_para(doc, "Adımlar:", bold=True)
    add_bullet(doc, "Drawer'da 'Element Lookup' butonuna tıkla (üstte mavi olur).")
    add_bullet(doc, "Altında bir terim input'u belirir.")
    add_bullet(doc, "Bir kelime yaz, örneğin: shaft, pad, gimbal, hub.")
    add_bullet(doc, "Ask butonuna bas.")

    add_para(doc, "Beklenen sonuç:", bold=True)
    add_bullet(doc, "Birkaç saniyede cevap döner.")
    add_bullet(doc, "Cevap metni şuna benzer: \"'shaft' was found in N places in the element patent analysis.\"")
    add_bullet(doc, "Altında Evidence kartları görünür; her kartın içinde arattığın kelime mavi highlight'lı olur.")
    add_bullet(doc, "Bulunamazsa: net 'not found' mesajı + alternatif terim önerisi gelir.")

    add_para(doc, "Olası sonuçlar:", bold=True)
    t = doc.add_table(rows=0, cols=2)
    t.style = "Light Grid"
    add_table_row(t, ["Sonuç", "Yorum"], header=True)
    add_table_row(t, ["Evidence kartları + highlight", "Backend, DB ve P3 mantığı çalışıyor."])
    add_table_row(t, ["'not found' mesajı", "P3 çalışıyor; sadece o terim bu patente yok. Başka kelime dene."])
    add_table_row(t, ["500 hata / 'Assistant call failed'", "Backend hatası. uvicorn terminalindeki traceback'e bak."])

    # TEST 2 — P1
    add_heading(doc, "TEST 2 — Core Problem (P1)", level=1, color=RGBColor(0x1F, 0x4E, 0x79))
    add_para(doc, "Neden ikinci bu: LLM çağrısı yapılan ilk test. Ollama kurulumunun "
             "sağlıklı çalıştığını doğrularsın.", italic=True)

    add_para(doc, "Adımlar:", bold=True)
    add_bullet(doc, "Drawer'da 'Core Problem' butonuna tıkla (zaten varsayılan, mavi).")
    add_bullet(doc, "Ask butonuna bas. (Term girmen gerekmez.)")
    add_bullet(doc, "ÖNEMLİ: İlk Ask çağrısı 30–90 saniye sürebilir. Ollama qwen2.5:7b modelini ilk kez RAM'e yüklerken yavaştır. İkinci çağrılar hızlı olur. Sabırlı ol.")

    add_para(doc, "Beklenen sonuç:", bold=True)
    add_bullet(doc, "Önce 'Thinking…' + 'Running model… this may take a few seconds' yazısı.")
    add_bullet(doc, "Sonra cevap döner: invention'ın çözmek istediği teknik problem ve çözüm yönü.")
    add_bullet(doc, "Üstte yeşil 'Explicitly Stated' rozeti VEYA kırmızı 'Insufficient' rozeti.")
    add_bullet(doc, "Evidence kartlarında prior_art_and_problems ve executive_summary alanlarından birebir alıntılar.")

    add_para(doc, "Olası sonuçlar:", bold=True)
    t = doc.add_table(rows=0, cols=2)
    t.style = "Light Grid"
    add_table_row(t, ["Sonuç", "Yorum"], header=True)
    add_table_row(t, ["Yeşil 'Explicitly Stated' + cevap + evidence", "Mükemmel. Modül tam çalışıyor."])
    add_table_row(t, ["Kırmızı 'Insufficient'", "Modül çalışıyor ama bu patentin prior_art_and_problems veya executive_summary alanı boş/yetersiz. Patent Inputs'tan doldur."])
    add_table_row(t, ["'Assistant call failed' / connection refused", "Ollama'ya bağlanamıyor. curl http://localhost:11434/api/tags ile kontrol et."])
    add_table_row(t, ["Çok uzun süre dönüyor (>3 dk)", "Model ilk yükleniyor olabilir veya makine yavaş. Bir kez daha bekle, sonra Ctrl+F5."])

    # TEST 3 — P2
    add_heading(doc, "TEST 3 — Claim Structure (P2)", level=1, color=RGBColor(0x2E, 0x7D, 0x32))
    add_para(doc, "Neden son bu: En kapsamlı modül. Hem LLM çağrısı hem yapılandırılmış "
             "çıktı (independent/dependent kart listeleri) hem de faithfulness "
             "değerlendirmesi yapılır.", italic=True)

    add_para(doc, "Adımlar:", bold=True)
    add_bullet(doc, "Drawer'da 'Claim Structure' butonuna tıkla.")
    add_bullet(doc, "Ask butonuna bas. (Term girmen gerekmez.)")

    add_para(doc, "Beklenen sonuç:", bold=True)
    add_bullet(doc, "Yine 30–90 saniyelik bekleme süresi.")
    add_bullet(doc, "Cevap içeriği:")
    add_bullet(doc, "    'Independent claim candidates' başlığı altında 1–2 kart (her kart: feature listesi + neden bu feature merkezi).")
    add_bullet(doc, "    'Dependent claim candidates' başlığı altında 0–N kart.")
    add_bullet(doc, "    'Drafting cautions' listesi (varsa).")
    add_bullet(doc, "    Her kartın altında sarı 'Inferred' veya yeşil 'Explicitly Stated' rozeti + support_note.")
    add_bullet(doc, "Evidence kartlarında executive_summary'den birebir alıntılar.")

    # Doğrulama
    add_heading(doc, "Doğrulama — uvicorn terminali", level=1)
    add_para(doc, "Her Ask basışında uvicorn terminalinde şu satıra benzer bir log çıkmalı:")
    add_code(doc, 'INFO:     127.0.0.1:xxxxx - "POST /api/patents/6/assistant/ask HTTP/1.1" 200 OK')

    add_para(doc, "Eğer 200 yerine 500 dönerse, hemen üstündeki traceback'i kopyala — hata burada.", italic=True)

    # Tarayıcı DevTools
    add_heading(doc, "Tarayıcı DevTools (F12) ile ekstra kontrol", level=1)
    add_bullet(doc, "Network sekmesi → 'assistant/ask' isteğini bul → Response'ta JSON dönmeli.")
    add_bullet(doc, "Console sekmesi → kırmızı hata yoksa frontend tamam.")

    # Sorun çıkarsa
    add_heading(doc, "Sorun çıkarsa kontrol noktaları", level=1)

    add_para(doc, "1) localhost:8000/api/health → şunu dönmeli:", bold=True)
    add_code(doc, '{"status":"ok"}')

    add_para(doc, "2) Ollama çalışıyor mu:", bold=True)
    add_code(doc, "curl http://localhost:11434/api/tags")
    add_para(doc, "Model listesi içinde qwen2.5:7b görmelisin.")

    add_para(doc, "3) Postgres çalışıyor mu:", bold=True)
    add_code(doc, "docker ps")
    add_para(doc, "paten_draft_backend-db-1 satırında 'Up X' ve 0.0.0.0:5433->5432/tcp olmalı.")

    add_para(doc, "4) İlk açılışta:", bold=True)
    add_bullet(doc, "Docker Desktop'ı manuel başlat, whale ikonu sabit olana kadar bekle.")
    add_bullet(doc, "PowerShell aç, cd C:\\Users\\melike\\paten_draft_backend.")
    add_bullet(doc, "docker-compose up -d (sadece DB için yeterli; app container'ı port çakışması olmasın diye DURDUR: docker stop paten_draft_backend-app-1).")
    add_bullet(doc, ".\\.venv\\Scripts\\Activate.ps1")
    add_bullet(doc, "uvicorn app.main:app --reload")
    add_bullet(doc, "Tarayıcıda localhost:8000 → patent seç → Open → drawer aç.")

    # Test sonucu kayıt formu
    add_heading(doc, "Test sonucu kayıt formu", level=1)
    add_para(doc, "Aşağıdaki tabloya test sonuçlarını yaz:")

    t = doc.add_table(rows=0, cols=3)
    t.style = "Light Grid"
    add_table_row(t, ["Test", "Sonuç (geçti / Insufficient / hata)", "Notlar"], header=True)
    add_table_row(t, ["P3 — Element Lookup", "", ""])
    add_table_row(t, ["P1 — Core Problem", "", ""])
    add_table_row(t, ["P2 — Claim Structure", "", ""])

    # Save
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
